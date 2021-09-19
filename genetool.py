import os
import sys
import doctest
import PySimpleGUI as sg
import docx
from docx import Document
from docx.enum.dml import MSO_COLOR_TYPE
from docx.shared import Inches
from docx.enum.text import WD_COLOR_INDEX


start = ["AUG", "ATG"]
stopU = ["UAA", "UAG", "UGA", "TAA", "TAG", "TGA"]
pheU = ["UUU", "UUC", "TTT", "TTC"]
leuU = ["UUA", "UUG", "TTA", "TTG"]
leuC = ["CUU", "CUC", "CUA", "CUG", "CTT", "CTC", "CTA", "CTG"]
ileA = ["AUU", "AUC", "AUA", "ATT", "ATC", "ATA"]
valG = ["GUU", "GUC", "GUA", "GUG", "GTT", "GTC", "GTA", "GTG"]
serU = ["UCU", "UCC", "UCA", "UCG", "TCT", "TCC", "TCA", "TCG"]
proC = ["CCU", "CCC", "CCA", "CCG", "CCT"]
thrA = ["ACU", "ACC", "ACA", "ACG", "ACT"]
alaG = ["GCU", "GCC", "GCA", "GCG", "GCT"]
tyrU = ["UAU", "UAC", "TAT", "TAC"]
hisC = ["CAU", "CAC", "CAT"]
glnC = ["CAA", "CAG"]
asnA = ["AAU", "AAC", "AAT"]
lysA = ["AAA", "AAG"]
aspG = ["GAU", "GAC", "GAT"]
gluG = ["GAA", "GAG"]
cysU = ["UGU", "UGC", "TGT", "TGC"]
trpU = ["UGG", "TGG"]
argC = ["CGU", "CGC", "CGA", "CGG", "CGT"]
glyG = ["GGU", "GGC", "GGA", "GGG", "GGT"]
serA = ["AGU", "AGC", "AGT"]
argA = ["AGA", "AGG"]

UUs = [pheU, leuU]
UCs = [serU]
UGs = [cysU, trpU, stopU]
UAs = [tyrU, stopU]

CUs = [leuC]
CCs = [proC]
CGs = [argC]
CAs = [hisC, glnC]

GUs = [valG]
GCs = [alaG]
GGs = [glyG]
GAs = [aspG, gluG]

AUs = [ileA, start]
ACs = [thrA]
AGs = [serA, argA]
AAs = [asnA, lysA]

AminoMap = {
    "*": stopU,
    "M": start,
    "F": pheU,
    "L": leuU,
    "l": leuC,
    "I": ileA,
    "V": valG,
    "S": serU,
    "P": proC,
    "T": thrA,
    "A": alaG,
    "Y": tyrU,
    "H": hisC,
    "Q": glnC,
    "N": asnA,
    "K": lysA,
    "D": aspG,
    "E": gluG,
    "C": cysU,
    "W": trpU,
    "R": argC,
    "G": glyG,
    "s": serA,
    "r": argA
}


# takes in a codon, searches through the Amino dict, and returns the amino acid letter ID
def codon_to_AA(codon):
    """tests with codons
    >>> codon_to_AA("AUG")
    'M'
    >>> codon_to_AA("ATG")
    'M'
    >>> codon_to_AA("UUU")
    'F'
    >>> codon_to_AA("UUC")
    'F'
    >>> codon_to_AA("TTT")
    'F'
    >>> codon_to_AA("TTC")
    'F'
    >>> codon_to_AA("UCU")
    'S'
    >>> codon_to_AA("UCC")
    'S'
    >>> codon_to_AA("CCU")
    'P'
    >>> codon_to_AA("CCG")
    'P'
    >>> codon_to_AA("ACA")
    'T'
    >>> codon_to_AA("ACU")
    'T'
    >>> codon_to_AA("UUA")
    'L'
    >>> codon_to_AA("CUC")
    'L'
    >>> codon_to_AA("AUU")
    'I'
    >>> codon_to_AA("AUA")
    'I'
    >>> codon_to_AA("GUC")
    'V'
    >>> codon_to_AA("GTG")
    'V'
    >>> codon_to_AA("GCG")
    'A'
    >>> codon_to_AA("GCA")
    'A'
    >>> codon_to_AA("UAU")
    'Y'
    >>> codon_to_AA("TAC")
    'Y'
    >>> codon_to_AA("CAT")
    'H'
    >>> codon_to_AA("CAC")
    'H'
    >>> codon_to_AA("CAA")
    'Q'
    >>> codon_to_AA("CAG")
    'Q'
    >>> codon_to_AA("AAT")
    'N'
    >>> codon_to_AA("AAC")
    'N'
    >>> codon_to_AA("GAT")
    'D'
    >>> codon_to_AA("GAC")
    'D'
    >>> codon_to_AA("GAA")
    'E'
    >>> codon_to_AA("GAG")
    'E'
    >>> codon_to_AA("TGT")
    'C'
    >>> codon_to_AA("UGC")
    'C'
    >>> codon_to_AA("UGC")
    'C'
    >>> codon_to_AA("TGG")
    'W'
    >>> codon_to_AA("CGT")
    'R'
    >>> codon_to_AA("CGG")
    'R'
    >>> codon_to_AA("AGU")
    'S'
    >>> codon_to_AA("AGC")
    'S'
    >>> codon_to_AA("AGA")
    'R'
    >>> codon_to_AA("AGG")
    'R'
    >>> codon_to_AA("GGT")
    'G'
    >>> codon_to_AA("GGA")
    'G'

    """
    curr = []
    if codon[0] == 'U' or codon[0] == 'T':
        if codon[1] == 'U' or codon[1] == 'T':
            curr = UUs
        elif codon[1] == 'C':
            curr = UCs
        elif codon[1] == 'A':
            curr = UAs
        elif codon[1] == 'G':
            curr = UGs

    elif codon[0] == 'C':
        if codon[1] == 'U' or codon[1] == 'T':
            curr = CUs
        elif codon[1] == 'C':
            curr = CCs
        elif codon[1] == 'A':
            curr = CAs
        elif codon[1] == 'G':
            curr = CGs

    elif codon[0] == 'A':
        if codon[1] == 'U' or codon[1] == 'T':
            curr = AUs
        elif codon[1] == 'C':
            curr = ACs
        elif codon[1] == 'A':
            curr = AAs
        elif codon[1] == 'G':
            curr = AGs

    elif codon[0] == 'G':
        if codon[1] == 'U' or codon[1] == 'T':
            curr = GUs
        elif codon[1] == 'C':
            curr = GCs
        elif codon[1] == 'A':
            curr = GAs
        elif codon[1] == 'G':
            curr = GGs

    for AA in curr:
        if codon in AA:
            amino = AA
    for key in AminoMap:
        if AminoMap[key] == amino:
            key = key.upper()
            return key


# takes in a string of nucleotides and outputs the amino acid sequence
def coding_to_AA(code):
    """
    >>> coding_to_AA("atggttcggaccgtcgcggtg")
    'MVRTVAV'
    >>> coding_to_AA("CGUCGGTGGAGGCGUUCGAGUGCU")
    'RRWRRSSA'
    >>> coding_to_AA(fileread("gene.txt"))
    'MVRRYLPLNPLRAFEAAARHLSFTRAAIELNVTHAAVSQQVRALEEQLGCVLFTRVSRGLVLTHEGEGLLPVLNEAFDRIADTLECFSHGQFRERVKVGAVGTFAAGWLLPRLAGFYDSHPHIDLHISTHNNHVDPAAEGHDYTIRFGNGAWHESDAELIFSAPHAPLCSPAIAEQLQQPDDVHRFTLLRSFRRDEWSRWLDCAGGTPPSPSQPVMVFDTSLAMAEAAQLGAGVAIAPVCMFSRLLQSGALVQPFAAEITLGGYWLTRLQSRTETPAMQQFARWLLNTAAA*'
    """
    codeU = ""
    for ch in code:
        codeU += ch.upper()

    out = ""
    i = 0
    while i + 3 <= len(codeU):
        cur = codeU[i: i + 3]
        amino = codon_to_AA(cur)
        out += amino
        i += 3
    return out


# takes in string and removes everything thats not ACTGU
def clean_string(s):
    """
    >>> clean_string("aacaa1234")
    'AACAA'
    >>> clean_string("acGtua1234")
    'ACGTUA'
    >>> clean_string(" a c   G  t u a 1 2 3 4 ")
    'ACGTUA'
    """

    out = ""
    for ch in s:
        ch = ch.upper()
        if ch.isalpha():
            if ch == 'T' or ch == 'U' or ch == 'A' or ch == 'C' or ch == 'G':
                out += ch
    return out


# takes in a filename, and outputs a single string of whats inside the file.
def fileread(filename):
    """
    >>> fileread ("gene.txt")
    'ATGGTCAGACGTTATCTCCCCCTTAACCCGCTGCGCGCCTTTGAGGCCGCCGCCCGTCATCTCAGTTTTACCCGCGCGGCGATTGAGCTGAATGTCACCCATGCCGCCGTCAGCCAGCAGGTCAGGGCGCTGGAAGAACAACTCGGCTGTGTGCTGTTTACCCGCGTCTCGCGCGGGCTGGTGCTGACCCATGAAGGTGAGGGATTACTGCCGGTGCTCAATGAGGCGTTTGACCGGATTGCGGATACTCTGGAGTGTTTTTCTCACGGGCAGTTCCGTGAGCGGGTGAAAGTCGGTGCGGTGGGAACATTTGCCGCAGGCTGGCTGCTGCCGCGTCTGGCCGGATTCTATGACAGCCATCCGCATATTGATCTGCATATCTCCACCCATAACAATCATGTGGACCCGGCGGCGGAAGGGCATGATTATACGATCCGTTTCGGTAACGGCGCGTGGCATGAGTCAGATGCGGAACTGATTTTCAGTGCACCACACGCTCCGCTGTGCTCACCGGCCATTGCAGAACAGTTACAGCAGCCGGATGATGTTCACCGCTTTACCCTGCTGCGCTCATTCCGCCGGGATGAATGGAGCCGCTGGCTGGATTGTGCGGGCGGCACACCGCCTTCCCCGTCACAGCCGGTAATGGTGTTCGATACCTCACTGGCCATGGCCGAGGCGGCACAACTGGGTGCCGGGGTAGCGATCGCACCGGTATGTATGTTCAGCCGCCTGTTACAGTCAGGCGCACTGGTACAGCCGTTTGCCGCAGAAATCACCCTCGGCGGCTACTGGCTGACGCGGTTACAGTCCCGTACGGAAACCCCGGCCATGCAGCAATTCGCCCGCTGGCTGCTGAATACGGCGGCGGCGTAA'
    """
    with open(filename, "r") as myfile:
        data = myfile.readlines()
    stripped = ""
    for elem in data:
        elem = elem.strip('\n')
        stripped += elem
    s = clean_string(stripped)  # stripped is now the full nucleotide sequence which is the product of this fx
    s2 = ""
    for ch in s:
        if ch == 'U':
            s2 += 'T'
        else:
            s2 += ch

    return s2


# takes in a gene, oligo, accaptable snps, and searches gene for matches that have less than the snp count
def search(gene, oligos, snps):
    """
    >>> search ('ATCGTTATCGTCGGTGGATC', ['GTTT', 'CAAA', 'TTTG', 'AAAC'], 1)
    [('GTTA', 1, 3, 'F')]
    >>> search ('ATCGTTATCGTCGGTGGATC', ['GGATC', 'CCTAG', 'CTAGG', 'GATCC'], 1)
    [('GGATC', 0, 15, 'F')]
    >>> search ('ATCGTTATCGTCGGTGGATC', oligo_complements('GTTATCGTCGG'), 0)
    [('GTTATCGTCGG', 0, 3, 'F')]
    """
    matches = []
    i = 0
    j = 0
    while j + len(oligos[0]) - 1 < len(gene):
        snpcountF = 0
        snpcountC = 0
        snpcountR = 0
        snpcountRC = 0
        while i < len(oligos[0]):
            query = j + len(oligos[0])
            cur1 = gene[j:query]
            if cur1[i] != oligos[0][i]:
                snpcountF += 1
            if cur1[i] != oligos[1][i]:
                snpcountC += 1
            if cur1[i] != oligos[2][i]:
                snpcountR += 1
            if cur1[i] != oligos[3][i]:
                snpcountRC += 1
            i += 1

        if snpcountF <= snps:
            tupF = (cur1, snpcountF, j, 'F')
            matches.append(tupF)
        if snpcountC <= snps:
            tupC = (cur1, snpcountC, j, 'C')
            matches.append(tupC)
        if snpcountR <= snps:
            tupR = (cur1, snpcountR, j, 'R')
            matches.append(tupR)
        if snpcountRC <= snps:
            tupRC = (cur1, snpcountRC, j, 'B')
            matches.append(tupRC)
        i = 0
        j += 1
    return matches


# takes in an oligo and returns a list of the oligo in forward , complement, reverse, reverse complement in that order
def oligo_complements(oligo):
    """
     >>> oligo_complements('AAAAA')
     ['AAAAA', 'TTTTT', 'AAAAA', 'TTTTT']
     >>> oligo_complements('CCA')
     ['CCA', 'GGT', 'ACC', 'TGG']
    """
    complement = ""
    reverse = ""
    reverse_complement = ""
    oligo2 = ""
    for ch in oligo:
        if ch == 'A':
            complement += 'T'
        elif ch == 'T' or ch == 'U':
            complement += 'A'
        elif ch == 'C':
            complement += 'G'
        elif ch == 'G':
            complement += 'C'
    for ch in oligo:
        if ch == 'U':
            oligo2 += 'T'
        else:
            oligo2 += ch

    i = len(oligo) - 1
    while i >= 0:
        reverse += oligo[i]
        reverse_complement += complement[i]
        i -= 1
    oligos = [oligo2, complement, reverse, reverse_complement]
    return oligos


def search_scrub(matches, searches):
    scrubbed_matches = []
    for match in matches:
        if match[3] == 'F' and searches[0]:
            scrubbed_matches.append(match)

        elif match[3] == 'C' and searches[1]:
            scrubbed_matches.append(match)

        elif match[3] == 'R' and searches[2]:
            scrubbed_matches.append(match)

        elif match[3] == 'B' and searches[3]:
            scrubbed_matches.append(match)

    return scrubbed_matches


# uses search results to print the highlighted file to a docx document and saves it.
def printout(oligos, gene, matches, outfilename, docx_opt, docx_indexes, docx_AA, aa_seq):
    match_indexes = {}
    for elem in matches:
        index = elem[2]
        typ = elem[3]
        match_indexes[index] = typ

    document = Document()
    document.add_paragraph()
    s = document.add_paragraph()
    s.add_run("your query oligo was: ").bold = True
    s.add_run(oligos[0])
    s = document.add_paragraph()
    s.add_run("complement oligo is: ").bold = True
    s.add_run(oligos[1])
    s = document.add_paragraph()
    s.add_run("reverse oligo is: ").bold = True
    s.add_run(oligos[2])
    s = document.add_paragraph()
    s.add_run("reverse complement oligo is: ").bold = True
    s.add_run(oligos[3])
    i = 0
    end = 0
    on = False
    oligol = len(matches[0][0])
    if docx_opt:
        s = document.add_paragraph()
        s.add_run("your highlighted gene is below. forward matches are green, complement matches in yellow, reverse matches in teal, and reverse complement matches in pink. ").bold = True
        p = document.add_paragraph()
        font = p.add_run().font
        while i < len(gene):
            if i in match_indexes or on:
                if i in match_indexes:
                    place = i
                    mat_typ = match_indexes[i]
                    end = place + oligol
                    on = True
                if on:
                    if mat_typ == 'F':
                        #p.highlight_color('GREEN')
                        p.add_run(gene[i]).font.highlight_color = WD_COLOR_INDEX.GREEN
                    elif mat_typ == 'C':
                        #p.highlight_color('YELLOW')
                        p.add_run(gene[i]).font.highlight_color = WD_COLOR_INDEX.YELLOW
                    elif mat_typ == 'R':
                        #p.highlight_color('TEAL')
                        p.add_run(gene[i]).font.highlight_color = WD_COLOR_INDEX.TEAL
                    elif mat_typ == 'B':
                        #p.highlight_color('PINK')
                        p.add_run(gene[i]).font.highlight_color = WD_COLOR_INDEX.PINK
            else:
                #p.highlight_color('AUTO')
                p.add_run(gene[i])

            if end == i:
                on = False
            i += 1
        document.add_paragraph()
        document.add_paragraph()
    if docx_AA:
        p = document.add_paragraph()
        p.add_run("(please note that amino acid conversion does not look for reading frame, and starts from the first 3 nucleotides)").bold=True
        p = document.add_paragraph()
        p.add_run("Amino Acid sequence below:").bold=True
        p = document.add_paragraph()
        p.add_run(aa_seq)

        document.add_paragraph()
        document.add_paragraph()
    if docx_indexes:
        p = document.add_paragraph()
        p.add_run("Index match info below:").bold=True
        for match in matches:
            p = document.add_paragraph()
            if match[3] == 'F':
                p.add_run("Forward sequence match found at index ")
                p.add_run(str(match[2]))
                p.add_run(" with ")
                p.add_run(str(match[1]))
                p.add_run(" SNPs")
                p = document.add_paragraph()
                p.add_run("Matching sequence : ")
                p.add_run(match[0])
                document.add_paragraph()
            elif match[3] == 'C':
                p.add_run("Complement sequence match found at index ")
                p.add_run(str(match[2]))
                p.add_run(" with ")
                p.add_run(str(match[1]))
                p.add_run(" SNPs")
                p = document.add_paragraph()
                p.add_run("Matching sequence : ")
                p.add_run(match[0])
                document.add_paragraph()
            elif match[3] == 'R':
                p.add_run("Reverse sequence match found at index ")
                p.add_run(str(match[2]))
                p.add_run(" with ")
                p.add_run(str(match[1]))
                p.add_run(" SNPs")
                p = document.add_paragraph()
                p.add_run("Matching sequence : ")
                p.add_run(match[0])
                document.add_paragraph()
            elif match[3] == 'B':
                p.add_run("Reverse complement sequence match found at index ")
                p.add_run(str(match[2]))
                p.add_run(" with ")
                p.add_run(str(match[1]))
                p.add_run(" SNPs")
                p = document.add_paragraph()
                p.add_run("Matching sequence : ")
                p.add_run(match[0])
                document.add_paragraph()
        document.add_paragraph()
        document.add_paragraph()

    document.save(outfilename)


def gui_ops():
    import PySimpleGUI as sg  # Part 1 - The import
    import colorama
    colorama.init()
    cprint = sg.cprint
    MLINE_KEY = '-ML-' + sg.WRITE_ONLY_KEY  # multiline element's key. Indicate it's an output only element
    output_key = MLINE_KEY
    # Define the window's contents
    layout = [[sg.Text("Please input name of .txt file to search or paste DNA sequence to search in text box")],
              [sg.Text("(gene highlights may not function beyond 1,000,000 base pairs)")],
              [sg.Input(), sg.Checkbox("check if filename in box")],
              [sg.Text("Please input the oligo you would like to search for")],
              [sg.Input()],
              [sg.Text("select the searches you would like to conduct ")],
              [sg.Checkbox("Forward search"), sg.Checkbox("Complement search"), sg.Checkbox("Reverse search"),
               sg.Checkbox("Reverse Compliment search")],
              [sg.Text("Please input the acceptable number of SNPs for your search")],
              [sg.Input()],
              [sg.Checkbox("print amino acid sequence of input gene?"),
               sg.Checkbox("print gene with highlights? (if unchecked will output indexes of matches)")],
              [sg.Text("Docx creation options below. Select all you would like included in your docx file")],
              [sg.Checkbox("create highlighted docx file of your gene?"), sg.Checkbox("include index results?"), sg.Checkbox("include amino acid sequence?")],
              [sg.Button('Ok'), sg.Button('close')],
              [sg.Multiline(size=(100, 40), key=MLINE_KEY)]
              ]

    # Create the window
    window = sg.Window('Welcome to Lazy Elephant III gene tools', layout)  # Part 3 - Window Defintion
    sg.cprint_set_output_destination(window, output_key)
    while True:
        event, values = window.read()  # Part 4 - Event loop or Window.read call
        if event == 'close' or event == sg.WIN_CLOSED:
            break

        if event == 'Ok':
            # variable assignment area
            if values[1]:
                filename = values[0]
                gene = fileread(filename)
            else:
                gene = values[0]
                gene = clean_string(gene)
            oligo = values[2]
            if oligo == '':
                oligo = 'A'
            oligo = oligo.upper()
            if values[7] != '' and values[7].isdigit():
                snps = int(values[7])
            elif values[7] == '':
                snps = 0
            else:
                cprint("snp value must be an integer", key=MLINE_KEY)
            printout_opt = False
            printout_opt = values[9]
            amino_convert_opt = values[8]

            # sub-function running area

            oligos = oligo_complements(oligo)
            aa_seq = coding_to_AA(gene)
            f_search = values[3]
            c_search = values[4]
            r_search = values[5]
            rc_search = values[6]
            docx_opt = values[10]
            docx_indexes = values[11]
            docx_AA = values[12]
            matches = search(gene, oligos, snps)
            searches = [f_search, c_search, r_search, rc_search]
            scrubbed = search_scrub(matches, searches)
            match_indexes = {}
            for elem in scrubbed:
                index = elem[2]
                typ = elem[3]
                match_indexes[index] = typ
            # console output area

            if printout_opt:
                i = 0
                o_len = len(oligo)
                on = False
                mat_typ = ""
                end_ind = 0
                cprint("(Forward sequence matches: green)", key=MLINE_KEY, end=' ')
                cprint("(Complement sequence matches: yellow)", key=MLINE_KEY, end=' ')
                cprint("(Reverse sequence matches: cyan)", key=MLINE_KEY, end=' ')
                cprint("(Reverse complement sequence matches: pink)", key=MLINE_KEY, end=' ')
                cprint(key=MLINE_KEY)
                cprint("Sequence of Interest", key=MLINE_KEY)
                while i < len(gene):
                    ch = gene[i]
                    if i in match_indexes:
                        end_ind = i + o_len
                        mat_typ = match_indexes[i]
                        on = True
                    if on:
                        if mat_typ == 'F':
                            cprint(ch, c=('black', 'green'), key=MLINE_KEY, end='')
                        if mat_typ == 'C':
                            cprint(ch, c=('black', 'yellow'), key=MLINE_KEY, end='')
                        if mat_typ == 'R':
                            cprint(ch, c=('black', 'cyan'), key=MLINE_KEY, end='')
                        if mat_typ == 'B':
                            cprint(ch, c=('black', 'pink'), key=MLINE_KEY, end='')
                    else:
                        cprint(ch, key=MLINE_KEY, end = '')
                    i += 1
                    if i == end_ind:
                        end_ind = 0
                        mat_typ = ''
                        on = False
                cprint(key=MLINE_KEY)
                cprint(key=MLINE_KEY)
            if amino_convert_opt:
                cprint("Amino acid sequence : ", key=MLINE_KEY)
                cprint(aa_seq, key=MLINE_KEY)
                cprint(key=MLINE_KEY)

            if not printout_opt and oligo != '':
                for match in scrubbed:
                    if match[3] == 'F':
                        cprint("Forward sequence match found at index ", key=MLINE_KEY, end='')
                        cprint(match[2], key=MLINE_KEY, end='')
                        cprint(" with ", key=MLINE_KEY, end='')
                        cprint(match[1], key=MLINE_KEY, end='')
                        cprint(" SNPs", key=MLINE_KEY)
                        cprint("Matching sequence : ", key=MLINE_KEY, end='')
                        cprint(match[0], key=MLINE_KEY)
                        cprint(key=MLINE_KEY)
                    elif match[3] == 'C':
                        cprint("Complement sequence match found at index ", key=MLINE_KEY, end='')
                        cprint(match[2], key=MLINE_KEY, end='')
                        cprint(" with ", key=MLINE_KEY, end='')
                        cprint(match[1], key=MLINE_KEY, end='')
                        cprint(" SNPs", key=MLINE_KEY)
                        cprint("Matching sequence : ", key=MLINE_KEY, end='')
                        cprint(match[0], key=MLINE_KEY)
                        cprint(key=MLINE_KEY)
                    elif match[3] == 'R':
                        cprint("Reverse sequence match found at index ", key=MLINE_KEY, end='')
                        cprint(match[2], key=MLINE_KEY, end='')
                        cprint(" with ", key=MLINE_KEY, end='')
                        cprint(match[1], key=MLINE_KEY, end='')
                        cprint(" SNPs", key=MLINE_KEY)
                        cprint("Matching sequence : ", key=MLINE_KEY, end='')
                        cprint(match[0], key=MLINE_KEY)
                        cprint(key=MLINE_KEY)
                    elif match[3] == 'B':
                        cprint("Reverse complement sequence match found at index ", key=MLINE_KEY, end='')
                        cprint(match[2], key=MLINE_KEY, end='')
                        cprint(" with ", key=MLINE_KEY, end='')
                        cprint(match[1], key=MLINE_KEY, end='')
                        cprint(" SNPs", key=MLINE_KEY)
                        cprint("Matching sequence : ", key=MLINE_KEY, end='')
                        cprint(match[0], key=MLINE_KEY)
                        cprint(key=MLINE_KEY)
            cprint(key=MLINE_KEY)
            cprint(key=MLINE_KEY)
            if docx_opt or docx_indexes or docx_AA:
                printout(oligos, gene, scrubbed, "gene_tool_results.docx", docx_opt, docx_indexes, docx_AA, aa_seq)

    # window.close()


def main():
    gui_ops()


main()
